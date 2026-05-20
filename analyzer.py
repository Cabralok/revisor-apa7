from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import requests
import os

GEMINI_KEY = os.environ.get("GEMINI_API_KEY", "")

FUENTES_PERMITIDAS = {
    "Times New Roman": 12,
    "Arial": 11,
    "Calibri": 11,
    "Georgia": 11,
    "Lucida Sans Unicode": 10,
}

def analizar_documento(ruta):
    doc = Document(ruta)
    
    dim1 = analizar_formato(doc)
    dim2, citas_encontradas = analizar_citas(doc)
    dim3, referencias_encontradas = analizar_referencias(doc)
    dim4 = analizar_coherencia(citas_encontradas, referencias_encontradas)
    dim5 = analizar_tablas(doc)
    dim6 = analizar_figuras(doc)
    dim7 = analizar_tabla_contenidos(doc)

    # Análisis IA para dimensiones 2, 3 y 4
    texto_completo = extraer_texto(doc)
    analisis_ia = analizar_con_ia(texto_completo, citas_encontradas, referencias_encontradas)
    
    # Merge IA results
    dim2["items"].extend(analisis_ia.get("citas_ia", []))
    dim3["items"].extend(analisis_ia.get("referencias_ia", []))
    dim4["items"].extend(analisis_ia.get("coherencia_ia", []))

    dimensiones = [dim1, dim2, dim3, dim4, dim5, dim6, dim7]
    
    # Calcular puntaje global
    total_items = sum(len(d["items"]) for d in dimensiones)
    items_ok = sum(
        sum(1 for i in d["items"] if i["estado"] == "ok")
        for d in dimensiones
    )
    puntaje = round((items_ok / total_items * 100)) if total_items > 0 else 0

    return {
        "puntaje": puntaje,
        "dimensiones": dimensiones,
        "nombre_archivo": os.path.basename(ruta)
    }


def extraer_texto(doc):
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def analizar_formato(doc):
    items = []

    # --- Fuente y tamaño ---
    fuentes_usadas = {}
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name and run.font.size:
                nombre = run.font.name
                tamanio = round(run.font.size.pt) if run.font.size else None
                fuentes_usadas[nombre] = tamanio

    if fuentes_usadas:
        fuente_principal = max(fuentes_usadas, key=lambda k: fuentes_usadas[k] or 0)
        tamanio_encontrado = fuentes_usadas[fuente_principal]
        tamanio_esperado = FUENTES_PERMITIDAS.get(fuente_principal)
        
        if fuente_principal in FUENTES_PERMITIDAS:
            if tamanio_encontrado == tamanio_esperado:
                items.append(item_ok("Fuente", f"{fuente_principal} {tamanio_encontrado}pt — correcto"))
            else:
                items.append(item_error(
                    "Fuente",
                    f"Se usa {fuente_principal} {tamanio_encontrado}pt",
                    f"Debe ser {fuente_principal} {tamanio_esperado}pt según APA 7"
                ))
        else:
            items.append(item_error(
                "Fuente",
                f"Fuente '{fuente_principal}' no está permitida",
                "APA 7 permite: Times New Roman 12, Arial 11, Calibri 11, Georgia 11, Lucida Sans Unicode 10"
            ))
    else:
        items.append(item_advertencia("Fuente", "No se pudo detectar la fuente", "Verificar manualmente"))

    # --- Márgenes ---
    section = doc.sections[0]
    margen_sup = round(section.top_margin.cm, 2) if section.top_margin else None
    margen_inf = round(section.bottom_margin.cm, 2) if section.bottom_margin else None
    margen_izq = round(section.left_margin.cm, 2) if section.left_margin else None
    margen_der = round(section.right_margin.cm, 2) if section.right_margin else None

    margenes_ok = all(
        m is not None and abs(m - 2.54) < 0.15
        for m in [margen_sup, margen_inf, margen_izq, margen_der]
    )

    if margenes_ok:
        items.append(item_ok("Márgenes", "2.54 cm en los cuatro lados — correcto"))
    else:
        detalle = f"Superior: {margen_sup}cm, Inferior: {margen_inf}cm, Izquierdo: {margen_izq}cm, Derecho: {margen_der}cm"
        items.append(item_error(
            "Márgenes",
            f"Márgenes incorrectos ({detalle})",
            "APA 7 requiere 2.54 cm en los cuatro lados"
        ))

    # --- Interlineado ---
    interlineados = []
    for p in doc.paragraphs:
        if p.paragraph_format.line_spacing:
            ls = p.paragraph_format.line_spacing
            if hasattr(ls, 'pt'):
                interlineados.append(round(ls.pt, 1))
            else:
                interlineados.append(float(ls))

    if interlineados:
        # Interlineado doble = 2.0 o aproximadamente 24pt
        dobles = sum(1 for x in interlineados if x == 2.0 or (20 <= x <= 28))
        porcentaje_doble = dobles / len(interlineados)
        if porcentaje_doble > 0.7:
            items.append(item_ok("Interlineado", "Doble en la mayoría del documento — correcto"))
        else:
            items.append(item_error(
                "Interlineado",
                "El interlineado no es doble en todo el documento",
                "APA 7 requiere interlineado doble en todo el documento"
            ))
    else:
        items.append(item_advertencia("Interlineado", "No se pudo verificar", "Revisar manualmente"))

    # --- Alineación ---
    alineaciones = []
    for p in doc.paragraphs:
        if p.text.strip() and p.paragraph_format.alignment is not None:
            alineaciones.append(p.paragraph_format.alignment)

    if alineaciones:
        izquierda = sum(1 for a in alineaciones if a == WD_ALIGN_PARAGRAPH.LEFT)
        justificado = sum(1 for a in alineaciones if a == WD_ALIGN_PARAGRAPH.JUSTIFY)
        if justificado > izquierda:
            items.append(item_error(
                "Alineación",
                "El texto está justificado",
                "APA 7 requiere alineación izquierda, no justificada"
            ))
        else:
            items.append(item_ok("Alineación", "Alineación izquierda — correcto"))
    else:
        items.append(item_advertencia("Alineación", "No se pudo verificar", "Revisar manualmente"))

    # --- Sangría ---
    sangrías_ok = 0
    sangrías_total = 0
    for p in doc.paragraphs:
        if p.text.strip() and len(p.text) > 50:
            sangrías_total += 1
            indent = p.paragraph_format.first_line_indent
            if indent and hasattr(indent, 'cm'):
                if abs(indent.cm - 1.27) < 0.2:
                    sangrías_ok += 1

    if sangrías_total > 0:
        if sangrías_ok / sangrías_total > 0.6:
            items.append(item_ok("Sangría de párrafo", "1.27 cm detectada — correcto"))
        else:
            items.append(item_error(
                "Sangría de párrafo",
                "La sangría de primera línea no es 1.27 cm",
                "APA 7 requiere sangría de 1.27 cm en la primera línea de cada párrafo"
            ))
    else:
        items.append(item_advertencia("Sangría de párrafo", "No se pudo verificar", "Revisar manualmente"))

    # --- Numeración de páginas ---
    tiene_numeracion = False
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header:
                for p in header.paragraphs:
                    if 'PAGE' in p._p.xml or any(run.text.strip().isdigit() for run in p.runs):
                        tiene_numeracion = True

    if tiene_numeracion:
        items.append(item_ok("Numeración de páginas", "Detectada — correcto"))
    else:
        items.append(item_advertencia(
            "Numeración de páginas",
            "No se detectó numeración de páginas",
            "APA 7 requiere numeración arábiga desde la primera página, en esquina superior derecha"
        ))

    # --- Portada ---
    primer_parrafo = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
    if primer_parrafo:
        items.append(item_advertencia(
            "Portada",
            "Verificar que el documento tenga portada con título, autor, institución",
            "APA 7 requiere portada con título (máx. 12 palabras), nombre del autor e institución"
        ))
    
    # --- Título en portada no supera 12 palabras ---
    if primer_parrafo:
        palabras = len(primer_parrafo.split())
        if palabras <= 12:
            items.append(item_ok("Longitud del título", f"{palabras} palabras — correcto"))
        else:
            items.append(item_error(
                "Longitud del título",
                f"El título tiene {palabras} palabras",
                "APA 7 recomienda que el título no supere 12 palabras"
            ))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Formato del documento", "puntaje": puntaje, "items": items}


def analizar_citas(doc):
    items = []
    texto = extraer_texto(doc)
    citas_encontradas = []

    # Patrón citas entre paréntesis: (Apellido, año) o (Apellido & Apellido, año)
    patron_parentetica = r'\(([A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ]+(?:\s+et\s+al\.)?(?:\s*[&y]\s*[A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ]+)?),\s*(\d{4}[a-z]?|s\.\s*f\.)\)'
    # Patrón cita narrativa: Apellido (año)
    patron_narrativa = r'([A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ]+(?:\s+et\s+al\.)?)\s+\((\d{4}[a-z]?|s\.\s*f\.)\)'

    citas_parenteticas = re.findall(patron_parentetica, texto)
    citas_narrativas = re.findall(patron_narrativa, texto)

    for autor, año in citas_parenteticas:
        citas_encontradas.append({"autor": autor.strip(), "año": año.strip(), "tipo": "parentetica"})
    for autor, año in citas_narrativas:
        citas_encontradas.append({"autor": autor.strip(), "año": año.strip(), "tipo": "narrativa"})

    if citas_encontradas:
        items.append(item_ok("Patrón de citas", f"Se detectaron {len(citas_encontradas)} citas en el texto"))
    else:
        items.append(item_advertencia(
            "Patrón de citas",
            "No se detectaron citas en el texto",
            "Verificar si el documento tiene citas en formato APA 7"
        ))

    # Verificar & vs y
    and_fuera = re.findall(r'[A-Z][a-z]+\s+&\s+[A-Z][a-z]+\s+\(', texto)
    y_dentro = re.findall(r'\([A-Z][a-z]+\s+y\s+[A-Z][a-z]+,\s*\d{4}', texto)

    if and_fuera:
        items.append(item_error(
            'Uso de "&" fuera del paréntesis',
            f'Se encontraron {len(and_fuera)} casos de "&" fuera del paréntesis',
            'Usar "y" cuando el autor va fuera del paréntesis. Ej: García y López (2020)'
        ))
    else:
        items.append(item_ok('Uso de "&" y "y"', 'No se detectaron errores de "&" fuera del paréntesis'))

    if y_dentro:
        items.append(item_error(
            'Uso de "y" dentro del paréntesis',
            f'Se encontraron {len(y_dentro)} casos de "y" dentro del paréntesis',
            'Usar "&" cuando los autores van dentro del paréntesis. Ej: (García & López, 2020)'
        ))
    else:
        items.append(item_ok('Uso de "&" dentro del paréntesis', 'No se detectaron errores'))

    # Verificar et al.
    et_al_mal = re.findall(r'et\.?\s*al[^.)]', texto)
    if et_al_mal:
        items.append(item_advertencia(
            'Formato de "et al."',
            f'Posibles errores en el formato de "et al." ({len(et_al_mal)} casos)',
            'El formato correcto es "et al." con punto y en cursiva'
        ))
    else:
        items.append(item_ok('Formato de "et al."', 'Formato correcto detectado'))

    # Citas largas en bloque (más de 40 palabras con comillas)
    comillas_largas = []
    for p in doc.paragraphs:
        if p.text.startswith('"') or p.text.startswith('"'):
            palabras = len(p.text.split())
            if palabras > 40:
                comillas_largas.append(p.text[:80])

    if comillas_largas:
        items.append(item_error(
            "Citas largas con comillas",
            f"Se encontraron {len(comillas_largas)} citas de más de 40 palabras entre comillas",
            "Las citas de 40 o más palabras deben ir en bloque con sangría, sin comillas"
        ))
    else:
        items.append(item_ok("Citas largas en bloque", "No se detectaron citas largas con comillas incorrectas"))

    # Citas sin número de página
    citas_con_pagina = re.findall(r'p\.\s*\d+|pp\.\s*\d+', texto)
    citas_textuales = re.findall(r'"[^"]{10,}"', texto)
    if citas_textuales and len(citas_con_pagina) < len(citas_textuales):
        items.append(item_advertencia(
            "Número de página en citas textuales",
            "Algunas citas textuales podrían no tener número de página",
            "Las citas directas (textuales) deben incluir número de página: (Autor, año, p. X)"
        ))
    else:
        items.append(item_ok("Número de página", "Las citas textuales detectadas incluyen número de página"))

    # s.f. para sin fecha
    sf_correcto = re.findall(r's\.\s*f\.', texto)
    sf_incorrecto = re.findall(r's/f|S/F|sf\b|S\.F\b', texto)
    if sf_incorrecto:
        items.append(item_error(
            'Formato "sin fecha"',
            f'Se encontró formato incorrecto para "sin fecha" ({len(sf_incorrecto)} casos)',
            'El formato correcto es "s. f." con espacios y puntos'
        ))
    else:
        items.append(item_ok('Formato "sin fecha"', 'No se detectaron errores en el formato de "s. f."'))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Citas en el texto", "puntaje": puntaje, "items": items}, citas_encontradas


def analizar_referencias(doc):
    items = []
    referencias_encontradas = []
    texto = extraer_texto(doc)

    # Detectar sección de referencias
    lineas = texto.split('\n')
    idx_referencias = -1
    for i, linea in enumerate(lineas):
        if re.match(r'^Referencias\s*$', linea.strip(), re.IGNORECASE):
            idx_referencias = i
            break

    if idx_referencias == -1:
        items.append(item_error(
            "Sección de referencias",
            "No se encontró la sección 'Referencias'",
            "El documento debe tener una sección titulada exactamente 'Referencias' al final"
        ))
        return {"nombre": "Lista de referencias", "puntaje": 0, "items": items}, []

    items.append(item_ok("Sección de referencias", "Se encontró la sección 'Referencias'"))

    # Verificar que "Referencias" esté centrado
    for p in doc.paragraphs:
        if p.text.strip().lower() == 'referencias':
            if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                items.append(item_ok("Título 'Referencias' centrado", "Correcto"))
            else:
                items.append(item_error(
                    "Título 'Referencias' centrado",
                    "El título 'Referencias' no está centrado",
                    "El título 'Referencias' debe estar centrado y sin negrita"
                ))
            # Verificar que no tenga negrita
            tiene_negrita = any(run.bold for run in p.runs)
            if tiene_negrita:
                items.append(item_error(
                    "Título 'Referencias' sin negrita",
                    "El título 'Referencias' tiene negrita",
                    "El título 'Referencias' debe estar centrado pero sin negrita en APA 7"
                ))
            else:
                items.append(item_ok("Título 'Referencias' sin negrita", "Correcto"))
            break

    # Extraer entradas de referencias
    bloque_refs = lineas[idx_referencias + 1:]
    entrada_actual = ""
    for linea in bloque_refs:
        if linea.strip():
            if re.match(r'^[A-ZÁÉÍÓÚÜÑ]', linea.strip()):
                if entrada_actual:
                    referencias_encontradas.append(entrada_actual.strip())
                entrada_actual = linea
            else:
                entrada_actual += " " + linea
    if entrada_actual:
        referencias_encontradas.append(entrada_actual.strip())

    if referencias_encontradas:
        items.append(item_ok("Referencias encontradas", f"Se detectaron {len(referencias_encontradas)} entradas en la lista de referencias"))
    else:
        items.append(item_advertencia("Referencias encontradas", "No se detectaron entradas en la lista de referencias", "Verificar formato"))

    # Verificar orden alfabético
    if len(referencias_encontradas) > 1:
        apellidos = []
        for ref in referencias_encontradas:
            match = re.match(r'^([A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ\-]+)', ref)
            if match:
                apellidos.append(match.group(1).lower())
        
        if apellidos == sorted(apellidos):
            items.append(item_ok("Orden alfabético", "Las referencias están en orden alfabético"))
        else:
            items.append(item_error(
                "Orden alfabético",
                "Las referencias no están en orden alfabético",
                "Ordenar alfabéticamente por el primer apellido del primer autor"
            ))

    # Verificar sangría francesa
    en_refs = False
    sangrías_francesas = 0
    total_refs_check = 0
    for p in doc.paragraphs:
        if p.text.strip().lower() == 'referencias':
            en_refs = True
            continue
        if en_refs and p.text.strip():
            total_refs_check += 1
            indent = p.paragraph_format.left_indent
            first = p.paragraph_format.first_line_indent
            if indent and first and hasattr(indent, 'cm') and hasattr(first, 'cm'):
                if indent.cm > 0 and first.cm < 0:
                    sangrías_francesas += 1

    if total_refs_check > 0:
        if sangrías_francesas / total_refs_check > 0.5:
            items.append(item_ok("Sangría francesa", "Detectada en la lista de referencias — correcto"))
        else:
            items.append(item_error(
                "Sangría francesa",
                "No se detectó sangría francesa en las referencias",
                "Cada entrada en la lista de referencias debe tener sangría francesa de 1.27 cm"
            ))

    # Verificar presencia de DOI o URL en referencias con artículos
    refs_con_doi = sum(1 for r in referencias_encontradas if 'doi.org' in r.lower() or 'https://' in r.lower())
    items.append(item_advertencia(
        "DOI o URL",
        f"{refs_con_doi} de {len(referencias_encontradas)} referencias tienen DOI o URL",
        "Los artículos de revista y recursos en línea deben incluir DOI o URL"
    ))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Lista de referencias", "puntaje": puntaje, "items": items}, referencias_encontradas


def analizar_coherencia(citas, referencias):
    items = []

    autores_citados = set()
    for c in citas:
        apellido = c["autor"].split()[0].lower()
        autores_citados.add((apellido, c["año"]))

    autores_referenciados = set()
    for ref in referencias:
        match = re.match(r'^([A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ\-]+).*?\((\d{4}[a-z]?|s\.\s*f\.)\)', ref)
        if match:
            autores_referenciados.add((match.group(1).lower(), match.group(2)))

    # Citas sin referencia
    sin_referencia = autores_citados - autores_referenciados
    if sin_referencia:
        for autor, año in sin_referencia:
            items.append(item_error(
                "Cita sin referencia",
                f'"{autor.capitalize()} ({año})" está citado en el texto pero no aparece en la lista de referencias',
                "Agregar la referencia completa en la lista de referencias"
            ))
    else:
        items.append(item_ok("Citas con referencia", "Todas las citas detectadas tienen su entrada en referencias"))

    # Referencias sin citar
    sin_citar = autores_referenciados - autores_citados
    if sin_citar:
        for autor, año in sin_citar:
            items.append(item_advertencia(
                "Referencia sin citar",
                f'"{autor.capitalize()} ({año})" aparece en referencias pero no fue citado en el texto',
                "Citar en el texto o eliminar de la lista de referencias"
            ))
    else:
        items.append(item_ok("Referencias citadas", "Todas las referencias detectadas fueron citadas en el texto"))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Coherencia citas ↔ referencias", "puntaje": puntaje, "items": items}


def analizar_tablas(doc):
    items = []
    tablas = doc.tables

    if not tablas:
        items.append(item_ok("Tablas", "No se encontraron tablas en el documento"))
        return {"nombre": "Tablas", "puntaje": 100, "items": items}

    items.append(item_ok("Tablas detectadas", f"Se encontraron {len(tablas)} tabla(s)"))

    for i, tabla in enumerate(tablas, 1):
        # Buscar párrafo anterior a la tabla
        # En python-docx, buscar el párrafo que precede a la tabla
        tabla_xml = tabla._tbl
        padre = tabla_xml.getparent()
        elementos = list(padre)
        idx = elementos.index(tabla_xml)
        
        parrafo_anterior = None
        for j in range(idx - 1, -1, -1):
            elem = elementos[j]
            if elem.tag.endswith('}p'):
                from docx.text.paragraph import Paragraph
                parrafo_anterior = Paragraph(elem, padre)
                break

        if parrafo_anterior:
            texto_titulo = parrafo_anterior.text.strip()
            
            # Verificar formato "Tabla N"
            if re.match(rf'^Tabla\s+{i}\b', texto_titulo):
                items.append(item_ok(f"Tabla {i}: numeración", "Formato 'Tabla N' correcto"))
            else:
                items.append(item_error(
                    f"Tabla {i}: numeración",
                    f"El título encontrado es '{texto_titulo[:50]}'",
                    f"Debe comenzar con 'Tabla {i}' sin punto ni símbolo"
                ))

            # Verificar cursiva en el título (segunda línea generalmente)
            tiene_cursiva = any(run.italic for run in parrafo_anterior.runs if run.text.strip())
            if not tiene_cursiva:
                items.append(item_error(
                    f"Tabla {i}: título en cursiva",
                    "El título descriptivo no está en cursiva",
                    "El título descriptivo de la tabla debe estar en cursiva"
                ))
            else:
                items.append(item_ok(f"Tabla {i}: título en cursiva", "Correcto"))

            # Sin negrita en título
            tiene_negrita = any(run.bold for run in parrafo_anterior.runs if run.text.strip())
            if tiene_negrita:
                items.append(item_error(
                    f"Tabla {i}: título sin negrita",
                    "El título tiene negrita",
                    "El título de la tabla no debe tener negrita"
                ))
            else:
                items.append(item_ok(f"Tabla {i}: título sin negrita", "Correcto"))
        else:
            items.append(item_error(
                f"Tabla {i}: título faltante",
                "No se encontró título encima de la tabla",
                f"Agregar 'Tabla {i}' seguido del título descriptivo en cursiva encima de la tabla"
            ))

        # Verificar nota debajo
        parrafo_siguiente = None
        for j in range(idx + 1, len(elementos)):
            elem = elementos[j]
            if elem.tag.endswith('}p'):
                from docx.text.paragraph import Paragraph
                parrafo_siguiente = Paragraph(elem, padre)
                break

        if parrafo_siguiente:
            texto_nota = parrafo_siguiente.text.strip()
            if texto_nota.startswith("Nota."):
                items.append(item_ok(f"Tabla {i}: nota", "Nota correctamente formateada con 'Nota.'"))
            elif texto_nota:
                items.append(item_advertencia(
                    f"Tabla {i}: nota",
                    "Hay texto debajo de la tabla que podría ser una nota mal formateada",
                    "Si es una nota, debe comenzar con 'Nota.' en cursiva"
                ))

        # Líneas verticales (no permitidas en APA 7)
        tiene_bordes_verticales = False
        for fila in tabla.rows:
            for celda in fila.cells:
                tc_pr = celda._tc.tcPr
                if tc_pr is not None:
                    bordes = tc_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left') + \
                             tc_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right')
                    for borde in bordes:
                        val = borde.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                        if val not in ['nil', 'none', '']:
                            tiene_bordes_verticales = True
                            break

        if tiene_bordes_verticales:
            items.append(item_error(
                f"Tabla {i}: líneas verticales",
                "La tabla tiene líneas verticales",
                "APA 7 no permite líneas verticales en las tablas. Solo líneas horizontales en encabezado y cierre"
            ))
        else:
            items.append(item_ok(f"Tabla {i}: sin líneas verticales", "Correcto"))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Tablas", "puntaje": puntaje, "items": items}


def analizar_figuras(doc):
    items = []
    
    # Detectar imágenes en el documento
    imagenes = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            imagenes.append(rel)

    if not imagenes:
        items.append(item_ok("Figuras", "No se encontraron figuras en el documento"))
        return {"nombre": "Figuras", "puntaje": 100, "items": items}

    items.append(item_advertencia(
        "Figuras detectadas",
        f"Se encontraron {len(imagenes)} imagen(es) en el documento",
        "Verificar que cada figura tenga su etiqueta 'Figura N' y título descriptivo en cursiva DEBAJO"
    ))

    # Buscar etiquetas "Figura N" en el texto
    texto = extraer_texto(doc)
    figuras_etiquetadas = re.findall(r'Figura\s+\d+', texto)

    if len(figuras_etiquetadas) < len(imagenes):
        items.append(item_error(
            "Etiquetas de figuras",
            f"Hay {len(imagenes)} imagen(es) pero solo {len(figuras_etiquetadas)} etiqueta(s) 'Figura N'",
            "Cada imagen debe tener 'Figura N' seguido del título descriptivo en cursiva DEBAJO de la imagen"
        ))
    else:
        items.append(item_ok("Etiquetas de figuras", f"{len(figuras_etiquetadas)} figura(s) etiquetadas correctamente"))

    # Verificar que el título va DEBAJO (en APA 7 va abajo, a diferencia de tablas)
    items.append(item_advertencia(
        "Ubicación del título de figuras",
        "Verificar que el título esté DEBAJO de cada figura",
        "A diferencia de las tablas, en figuras el título va DEBAJO de la imagen en APA 7"
    ))

    # Verificar numeración sin punto
    figuras_con_punto = re.findall(r'Figura\s+\d+\.', texto)
    if figuras_con_punto:
        items.append(item_error(
            "Numeración de figuras",
            f"Se encontraron {len(figuras_con_punto)} figuras con punto al final del número",
            "Formato correcto: 'Figura 1' sin punto. El punto va al final del título descriptivo"
        ))
    else:
        items.append(item_ok("Numeración de figuras", "Sin punto al final del número — correcto"))

    # Referencias a figuras en el texto
    figuras_citadas = re.findall(r'(?:ver\s+|véase\s+|en\s+la\s+)?[Ff]igura\s+\d+', texto)
    if figuras_citadas:
        items.append(item_ok("Figuras citadas en el texto", f"{len(figuras_citadas)} referencia(s) a figuras en el texto"))
    else:
        items.append(item_advertencia(
            "Figuras citadas en el texto",
            "No se detectaron referencias a las figuras en el cuerpo del texto",
            "Cada figura debe ser citada en el texto antes de aparecer. Ej: '(ver Figura 1)'"
        ))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Figuras", "puntaje": puntaje, "items": items}


def analizar_tabla_contenidos(doc):
    items = []
    texto = extraer_texto(doc)

    # Detectar si hay tabla de contenidos
    patrones_toc = [
        r'[Tt]abla\s+de\s+[Cc]ontenidos?',
        r'[Íí]ndice\s+(?:general|de\s+contenidos?)?',
        r'[Cc]ontenidos?'
    ]
    
    tiene_toc = any(re.search(p, texto) for p in patrones_toc)

    if not tiene_toc:
        items.append(item_advertencia(
            "Tabla de contenidos",
            "No se detectó tabla de contenidos",
            "Si el trabajo lo requiere, incluir tabla de contenidos con formato APA 7"
        ))
        return {"nombre": "Tabla de contenidos", "puntaje": 50, "items": items}

    items.append(item_ok("Tabla de contenidos", "Se detectó tabla de contenidos"))

    # Verificar título centrado
    for p in doc.paragraphs:
        texto_p = p.text.strip().lower()
        if any(re.match(pat.lower(), texto_p) for pat in ['tabla de contenidos?', 'índice', 'contenidos?']):
            if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                items.append(item_ok("Título del índice centrado", "Correcto"))
            else:
                items.append(item_error(
                    "Título del índice centrado",
                    "El título de la tabla de contenidos no está centrado",
                    "El título debe estar centrado"
                ))
            break

    # Verificar que la portada no aparece como entrada
    if re.search(r'[Pp]ortada\s*\.{3,}\s*\d', texto):
        items.append(item_error(
            "Portada en el índice",
            "La portada aparece como entrada en la tabla de contenidos",
            "La portada no debe incluirse como entrada en la tabla de contenidos"
        ))
    else:
        items.append(item_ok("Portada no incluida en el índice", "Correcto"))

    # Verificar que la propia tabla de contenidos no aparece como entrada
    if re.search(r'[Tt]abla\s+de\s+[Cc]ontenidos?\s*\.{3,}\s*\d', texto):
        items.append(item_error(
            "Tabla de contenidos en el índice",
            "La tabla de contenidos se incluye a sí misma como entrada",
            "La tabla de contenidos no debe aparecer como entrada de sí misma"
        ))
    else:
        items.append(item_ok("Tabla de contenidos no se auto-referencia", "Correcto"))

    # Verificar puntos de relleno
    tiene_puntos_relleno = bool(re.search(r'\w\s*\.{3,}\s*\d', texto))
    if tiene_puntos_relleno:
        items.append(item_ok("Puntos de relleno", "Se detectaron puntos de relleno entre entradas y números de página"))
    else:
        items.append(item_advertencia(
            "Puntos de relleno",
            "No se detectaron puntos de relleno entre entradas y números de página",
            "Usar puntos de relleno (....) entre el título de la sección y el número de página"
        ))

    puntaje = calcular_puntaje(items)
    return {"nombre": "Tabla de contenidos", "puntaje": puntaje, "items": items}


def analizar_con_ia(texto, citas, referencias):
    """Usa Gemini para análisis semántico profundo"""
    if not GEMINI_KEY or not texto:
        return {}

    prompt = f"""Sos un experto en Normas APA 7ª edición. Analizá el siguiente fragmento de un trabajo académico y detectá errores específicos.

TEXTO DEL DOCUMENTO (primeros 3000 caracteres):
{texto[:3000]}

CITAS DETECTADAS: {[f"{c['autor']} ({c['año']})" for c in citas[:10]]}
REFERENCIAS DETECTADAS: {referencias[:5]}

Respondé SOLO en JSON con esta estructura exacta:
{{
  "citas_ia": [
    {{"estado": "error|advertencia|ok", "titulo": "...", "detalle": "...", "sugerencia": "..."}}
  ],
  "referencias_ia": [
    {{"estado": "error|advertencia|ok", "titulo": "...", "detalle": "...", "sugerencia": "..."}}
  ],
  "coherencia_ia": [
    {{"estado": "error|advertencia|ok", "titulo": "...", "detalle": "...", "sugerencia": "..."}}
  ]
}}

Solo incluí ítems donde detectes algo relevante. Máximo 5 ítems por categoría. Respondé SOLO el JSON, sin texto adicional."""

    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_KEY}"
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        response = requests.post(url, json=payload, timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            texto_respuesta = data['candidates'][0]['content']['parts'][0]['text']
            texto_limpio = texto_respuesta.strip().replace('```json', '').replace('```', '').strip()
            import json
            resultado = json.loads(texto_limpio)
            
            # Normalizar formato
            for categoria in ['citas_ia', 'referencias_ia', 'coherencia_ia']:
                if categoria in resultado:
                    items_normalizados = []
                    for it in resultado[categoria]:
                        items_normalizados.append(item_desde_ia(it))
                    resultado[categoria] = items_normalizados
            return resultado
    except Exception as e:
        print(f"Error IA: {e}")
    
    return {}


def item_desde_ia(it):
    estado = it.get("estado", "advertencia")
    if estado == "error":
        return item_error(it.get("titulo", ""), it.get("detalle", ""), it.get("sugerencia", ""))
    elif estado == "ok":
        return item_ok(it.get("titulo", ""), it.get("detalle", ""))
    else:
        return item_advertencia(it.get("titulo", ""), it.get("detalle", ""), it.get("sugerencia", ""))


# --- Helpers ---

def item_ok(titulo, detalle):
    return {"estado": "ok", "titulo": titulo, "detalle": detalle, "sugerencia": ""}

def item_error(titulo, detalle, sugerencia):
    return {"estado": "error", "titulo": titulo, "detalle": detalle, "sugerencia": sugerencia}

def item_advertencia(titulo, detalle, sugerencia):
    return {"estado": "advertencia", "titulo": titulo, "detalle": detalle, "sugerencia": sugerencia}

def calcular_puntaje(items):
    if not items:
        return 0
    ok = sum(1 for i in items if i["estado"] == "ok")
    advertencia = sum(1 for i in items if i["estado"] == "advertencia")
    total = len(items)
    return round((ok + advertencia * 0.5) / total * 100)
